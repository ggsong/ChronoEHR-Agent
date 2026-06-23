#!/usr/bin/env python3
"""Export ChronoEHR manuscript drafts as DOCX files.

This script converts completed ChronoEHR-Agent outputs into Word documents.
It does not read raw EHR data and does not rerun any model.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Iterable

import pandas as pd

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - only exercised in missing dependency environments.
    raise SystemExit(
        "python-docx is required. Install requirements.txt or run this command with a Python "
        "environment that has python-docx available."
    ) from exc

DEFAULT_PROJECT = Path(__file__).resolve().parents[2]
REPORT_DIR = Path("outputs/reports")
SUPPLEMENT_DIR = Path("outputs/tables/supplementary_appendix")
TABLE_DIR = Path("outputs/tables")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "B8C2CC"

DEFAULT_EXPORT_CONFIG: dict[str, Any] = {
    "language": "zh",
    "outputs": {
        "main_docx": str(REPORT_DIR / "ChronoEHR_Methods_Results_Draft.docx"),
        "supplement_docx": str(REPORT_DIR / "ChronoEHR_Supplementary_Appendix.docx"),
    },
    "main_document": {
        "include_sections": {
            "title": True,
            "methods": True,
            "cohort_summary": True,
            "prediction_time_delta": True,
            "feature_group_ablation": True,
            "selected_feature_sets": True,
            "ed_los_sensitivity": True,
            "threshold_analysis": True,
            "decision_curve": True,
            "subgroup_performance": True,
            "next_writing_suggestions": True,
        }
    },
    "supplement": {
        "include_tables": ["S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8", "S9", "S10"],
        "max_rows": {"S4": 30, "S10": 90},
    },
}


def deep_merge(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    merged = dict(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = deep_merge(merged[key], value)
        else:
            merged[key] = value
    return merged


def load_export_config(path: Path | None) -> dict[str, Any]:
    if path is None:
        return DEFAULT_EXPORT_CONFIG
    if path.suffix.lower() not in {".json"}:
        raise SystemExit(
            "Manuscript export config currently uses JSON to avoid extra dependencies. "
            "Use configs/manuscript_export_template.json as the executable config."
        )
    if not path.exists():
        raise FileNotFoundError(f"Missing manuscript export config: {path}")
    user_config = json.loads(path.read_text(encoding="utf-8"))
    return deep_merge(DEFAULT_EXPORT_CONFIG, user_config)


def main_section_enabled(config: dict[str, Any], section: str) -> bool:
    sections = config.get("main_document", {}).get("include_sections", {})
    return bool(sections.get(section, True))


def supplement_table_enabled(config: dict[str, Any], table_id: str) -> bool:
    include_tables = config.get("supplement", {}).get("include_tables", DEFAULT_EXPORT_CONFIG["supplement"]["include_tables"])
    return table_id in set(map(str, include_tables))


def configured_max_rows(config: dict[str, Any], table_id: str, default: int | None) -> int | None:
    max_rows = config.get("supplement", {}).get("max_rows", {})
    value = max_rows.get(table_id, default)
    return None if value is None else int(value)


def configured_supplement_label(config: dict[str, Any]) -> str:
    include_tables = config.get("supplement", {}).get("include_tables", DEFAULT_EXPORT_CONFIG["supplement"]["include_tables"])
    labels = [str(value) for value in include_tables]
    if labels == DEFAULT_EXPORT_CONFIG["supplement"]["include_tables"]:
        return "完整 S1-S10"
    return "、".join(labels)


def read_csv(project_root: Path, relative_path: Path) -> pd.DataFrame:
    path = project_root / relative_path
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def fmt_value(value: object) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        return f"{value:.3f}"
    return str(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def remove_style_paragraph_borders(style) -> None:
    p_pr = style._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def set_table_width(table, total_width_dxa: int, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_width_dxa))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is not None:
        for col in list(grid):
            grid.remove(col)
    else:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def style_cell_text(cell, *, bold: bool = False, size: float = 8.0, color: RGBColor | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold
            if color is not None:
                run.font.color.rgb = color


def apply_compact_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.color.rgb = DARK_BLUE
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.15
    remove_style_paragraph_borders(title)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def configure_page(document: Document, *, landscape: bool = False) -> None:
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        margin = Inches(0.55)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        margin = Inches(1)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_footer(document: Document, label: str) -> None:
    section = document.sections[0]
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = label
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = MUTED


def add_note_table(document: Document, rows: Iterable[tuple[str, str]], *, width_dxa: int = 9360) -> None:
    table = document.add_table(rows=0, cols=2)
    set_table_borders(table)
    set_table_width(table, width_dxa, [1800, width_dxa - 1800])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_FILL)
        style_cell_text(cells[0], bold=True, size=9)
        style_cell_text(cells[1], size=9)


def add_dataframe_table(
    document: Document,
    dataframe: pd.DataFrame,
    columns: list[str],
    display_names: dict[str, str] | None = None,
    *,
    width_dxa: int = 9360,
    max_rows: int | None = None,
    font_size: float = 8.0,
) -> None:
    display_names = display_names or {}
    df = dataframe.copy()
    if max_rows is not None:
        df = df.head(max_rows)
    table = document.add_table(rows=1, cols=len(columns))
    set_table_borders(table)
    equal = width_dxa // len(columns)
    widths = [equal for _ in columns]
    widths[-1] = width_dxa - sum(widths[:-1])
    set_table_width(table, width_dxa, widths)

    for idx, column in enumerate(columns):
        cell = table.rows[0].cells[idx]
        cell.text = display_names.get(column, column)
        set_cell_shading(cell, HEADER_FILL)
        style_cell_text(cell, bold=True, size=font_size, color=DARK_BLUE)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, column in enumerate(columns):
            cells[idx].text = fmt_value(row.get(column, ""))
            style_cell_text(cells[idx], size=font_size)


def prepare_cohort_table(project_root: Path) -> pd.DataFrame:
    df = read_csv(project_root, TABLE_DIR / "chronic_disease_benchmark_cohort_summary.csv")
    if df.empty:
        return df
    out = df.copy()
    out["readmission_30d_rate_pct"] = out["readmission_30d_rate"].map(lambda value: f"{value:.2%}")
    return out.rename(
        columns={
            "cohort": "队列",
            "final_index_admissions": "住院次数",
            "final_subjects": "患者数",
            "readmission_30d_count": "30天再入院数",
            "readmission_30d_rate_pct": "30天再入院率",
        }
    )


def prepare_prediction_delta(project_root: Path) -> pd.DataFrame:
    df = read_csv(project_root, TABLE_DIR / "chronic_disease_prediction_time_deltas.csv")
    if df.empty:
        return df
    return df.rename(
        columns={
            "cohort_label": "队列",
            "admission_AUROC": "入院AUROC",
            "discharge_AUROC": "出院AUROC",
            "delta_AUROC": "AUROC差值",
            "admission_AUPRC": "入院AUPRC",
            "discharge_AUPRC": "出院AUPRC",
            "delta_AUPRC": "AUPRC差值",
        }
    )


def add_main_document(project_root: Path, output: Path, config: dict[str, Any]) -> None:
    document = Document()
    apply_compact_styles(document)
    configure_page(document)
    add_footer(document, "ChronoEHR-Agent | Methods/Results Draft")

    if main_section_enabled(config, "title"):
        title = document.add_paragraph(style="Title")
        title.add_run("ChronoEHR-Agent 慢病 EHR Time-Aware Benchmark")
        subtitle = document.add_paragraph("Methods/Results 草稿与论文写作材料包")
        subtitle.paragraph_format.space_after = Pt(8)
        for run in subtitle.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.font.color.rgb = MUTED

        document.add_paragraph(
            "本文件由本地 ChronoEHR-Agent 输出自动整理生成，只用于 EHR 数据研究、模型评估和论文草稿准备；"
            "不提供临床诊断、治疗或用药建议。"
        )
        add_note_table(
            document,
            [
                ("Design preset", "compact_reference_guide；主文稿使用 US Letter portrait，表格采用 9360 DXA 固定宽度。"),
                ("Data source", "本地 MIMIC-IV 3.1；慢病队列包括糖尿病、CKD、心衰和高血压。"),
                ("Primary task", "30-day all-cause readmission prediction with admission, inhospital 24h, and discharge prediction times."),
                ("Model boundary", "传统 baseline 为 logistic regression、Random Forest、HistGradientBoosting 和校准模型；LLM/Agent 不替代模型本身。"),
            ],
        )

        document.add_heading("建议题目", level=1)
        document.add_paragraph("中文：预测时间点和特征可用性对慢病住院患者30天再入院预测性能的影响：一项基于 MIMIC-IV 的 time-aware EHR benchmark。")
        document.add_paragraph("English: Impact of prediction time and feature availability on 30-day readmission prediction in chronic disease cohorts: a time-aware EHR benchmark using MIMIC-IV.")

    if main_section_enabled(config, "methods"):
        document.add_heading("Methods 草稿", level=1)
        for heading, text in [
            (
                "数据来源与工具边界",
                "本研究使用本地 MIMIC-IV 3.1 数据库开展回顾性 EHR 数据分析。ChronoEHR-Agent 被设计为本地研究工具，"
                "用于队列构建、预测时间点定义、时间窗内特征抽取、特征泄漏审计、传统机器学习 baseline 建模和结果报告生成。"
                "工具不输出临床诊断或治疗建议。",
            ),
            (
                "研究队列",
                "本轮 benchmark 包含四个慢病相关住院队列：糖尿病、CKD、心衰和高血压。每个队列均基于诊断编码识别研究对象，"
                "并以住院记录作为 index admission。模型评估采用患者级切分，避免同一患者同时出现在训练集和测试集。",
            ),
            (
                "结局与预测时间点",
                "主要结局为出院后30天内 all-cause hospital readmission。预测时间点包括入院时、入院后24小时和出院前。"
                "核心原则是只允许使用预测时已经可见的信息，不能让出院后或随访窗口内才产生的信息进入特征。",
            ),
            (
                "模型与指标",
                "主要 baseline 为 logistic regression，并与 Random Forest、HistGradientBoosting、Platt scaling 和 isotonic calibration 对照。"
                "报告 AUROC、AUPRC、Brier score 和 calibration decile summaries；类别不平衡场景下不只看 AUROC。",
            ),
        ]:
            document.add_heading(heading, level=2)
            document.add_paragraph(text)

    document.add_heading("Results 摘要", level=1)
    cohort = prepare_cohort_table(project_root)
    if main_section_enabled(config, "cohort_summary") and not cohort.empty:
        add_caption(document, "Table 1. Cohort summary.")
        add_dataframe_table(
            document,
            cohort,
            ["队列", "住院次数", "患者数", "30天再入院数", "30天再入院率"],
            width_dxa=9360,
            font_size=8.5,
        )

    delta = prepare_prediction_delta(project_root)
    if main_section_enabled(config, "prediction_time_delta") and not delta.empty:
        document.add_heading("Prediction time 对模型表现的影响", level=2)
        document.add_paragraph(
            "从入院时到出院前，四个慢病队列的 AUROC 和 AUPRC 均有不同程度提升，说明预测时间点本身是 EHR 模型结果解释的重要部分。"
        )
        add_caption(document, "Table 2. Admission-to-discharge performance delta.")
        add_dataframe_table(
            document,
            delta,
            ["队列", "入院AUROC", "出院AUROC", "AUROC差值", "入院AUPRC", "出院AUPRC", "AUPRC差值"],
            width_dxa=9360,
            font_size=7.5,
        )

    ablation = read_csv(project_root, SUPPLEMENT_DIR / "table_s3_feature_group_ablation.csv")
    if main_section_enabled(config, "feature_group_ablation") and not ablation.empty:
        document.add_heading("Feature group ablation", level=2)
        document.add_paragraph(
            "Grouped ablation 显示，24h labs、discharge labs 和 broad medication features 是较稳定的增量来源；"
            "ICU vitals/procedures 的平均增益较小且方向不稳定。"
        )
        add_caption(document, "Table 3. Feature group ablation summary.")
        add_dataframe_table(
            document,
            ablation,
            [
                "stage",
                "group_added",
                "comparisons",
                "cohorts_improved_AUROC",
                "cohorts_improved_AUPRC",
                "mean_delta_AUROC",
                "mean_delta_AUPRC",
                "mean_delta_Brier",
            ],
            width_dxa=9360,
            font_size=6.8,
        )

    selected = read_csv(project_root, SUPPLEMENT_DIR / "table_s5_selected_feature_set_comparison.csv")
    if main_section_enabled(config, "selected_feature_sets") and not selected.empty:
        document.add_heading("Selected feature set 对照", level=2)
        document.add_paragraph(
            "基于 repeated concepts 构造的 selected feature sets 保留了大部分性能，同时变量数量更少、解释更清楚，适合作为论文中的简化模型或敏感性分析。"
        )
        compact_selected = selected[
            [
                "cohort_label",
                "prediction_time",
                "selected_features",
                "full_AUROC",
                "selected_AUROC",
                "delta_AUROC",
                "full_AUPRC",
                "selected_AUPRC",
                "delta_AUPRC",
                "mean_absolute_calibration_error",
            ]
        ].copy()
        add_caption(document, "Table 4. Full vs selected feature set comparison.")
        add_dataframe_table(
            document,
            compact_selected,
            list(compact_selected.columns),
            width_dxa=9360,
            font_size=6.5,
        )

    ed_los = read_csv(project_root, SUPPLEMENT_DIR / "table_s7_ed_los_sensitivity.csv")
    if main_section_enabled(config, "ed_los_sensitivity") and not ed_los.empty:
        document.add_heading("ED length-of-stay 敏感性分析", level=2)
        mean_abs_auroc = ed_los["delta_AUROC"].abs().mean()
        mean_abs_auprc = ed_los["delta_AUPRC"].abs().mean()
        document.add_paragraph(
            f"Leakage gate 将 `ed_los_hours` 标记为 24h prediction 的 conditional availability 变量。"
            f"去掉该变量后，17 个 24h logistic models 的平均绝对 AUROC 变化为 {mean_abs_auroc:.4f}，"
            f"平均绝对 AUPRC 变化为 {mean_abs_auprc:.4f}，说明主要 24h 结论不依赖这个边界时间变量。"
        )
        ed_los_summary = (
            ed_los.groupby("cohort_label", sort=False)
            .agg(
                comparisons=("source_feature_set", "count"),
                mean_delta_AUROC=("delta_AUROC", "mean"),
                mean_delta_AUPRC=("delta_AUPRC", "mean"),
                mean_delta_Brier=("delta_Brier", "mean"),
                max_abs_delta_AUROC=("delta_AUROC", lambda values: float(values.abs().max())),
                max_abs_delta_AUPRC=("delta_AUPRC", lambda values: float(values.abs().max())),
            )
            .reset_index()
        )
        add_caption(document, "Table 5. ED length-of-stay sensitivity summary.")
        add_dataframe_table(
            document,
            ed_los_summary,
            list(ed_los_summary.columns),
            width_dxa=9360,
            font_size=7.2,
        )

    threshold = read_csv(project_root, SUPPLEMENT_DIR / "table_s8_threshold_analysis.csv")
    if main_section_enabled(config, "threshold_analysis") and not threshold.empty:
        document.add_heading("固定 alert burden 阈值分析", level=2)
        top10 = threshold[threshold["alert_rate"].eq(0.10)].copy()
        mean_ppv = top10["ppv"].mean()
        mean_recall = top10["recall"].mean()
        document.add_paragraph(
            f"固定 alert burden 用来补充 AUROC/AUPRC。Top 10% alert burden 下，最终 24h 和出院前模型的"
            f"平均 PPV 为 {mean_ppv:.3f}，平均 recall 为 {mean_recall:.3f}。该结果只描述模型排序行为，不是临床处置建议。"
        )
        top10_compact = top10[
            [
                "cohort_label",
                "prediction_time",
                "alerts",
                "ppv",
                "recall",
                "specificity",
                "lift_vs_event_rate",
            ]
        ].copy()
        add_caption(document, "Table 6. Top 10% alert burden summary.")
        add_dataframe_table(
            document,
            top10_compact,
            list(top10_compact.columns),
            width_dxa=9360,
            font_size=7.2,
        )

    decision_curve = read_csv(project_root, SUPPLEMENT_DIR / "table_s9_decision_curve.csv")
    if main_section_enabled(config, "decision_curve") and not decision_curve.empty:
        document.add_heading("Decision-curve net benefit 分析", level=2)
        threshold20 = decision_curve[decision_curve["threshold_probability"].eq(0.20)].copy()
        model_preferred = int((threshold20["preferred_strategy"] == "model").sum())
        mean_advantage = threshold20["net_benefit_advantage"].mean()
        document.add_paragraph(
            f"Decision-curve analysis 用来检查模型在给定研究风险阈值下是否优于 treat-all 和 treat-none 参照策略。"
            f"在 threshold probability = 0.20 时，{model_preferred}/{len(threshold20)} 个最终模型的 net benefit 更高，"
            f"平均 net benefit advantage 为 {mean_advantage:.4f}。该分析是模型评估，不定义临床处置阈值。"
        )
        decision_compact = threshold20[
            [
                "cohort_label",
                "prediction_time",
                "alerts",
                "alert_rate",
                "ppv",
                "recall",
                "model_net_benefit",
                "net_benefit_advantage",
                "preferred_strategy",
            ]
        ].copy()
        add_caption(document, "Table 7. Decision-curve summary at threshold probability 0.20.")
        add_dataframe_table(
            document,
            decision_compact,
            list(decision_compact.columns),
            width_dxa=9360,
            font_size=6.8,
        )

    subgroup = read_csv(project_root, SUPPLEMENT_DIR / "table_s10_subgroup_performance.csv")
    if main_section_enabled(config, "subgroup_performance") and not subgroup.empty:
        document.add_heading("Subgroup performance 分析", level=2)
        subgroup_summary = (
            subgroup.groupby("subgroup_variable", sort=False)
            .agg(
                rows=("subgroup_value", "count"),
                min_n=("n", "min"),
                mean_event_rate=("event_rate", "mean"),
                mean_AUROC=("AUROC", "mean"),
                mean_AUPRC=("AUPRC", "mean"),
                mean_top10_ppv=("top10_ppv", "mean"),
            )
            .reset_index()
        )
        document.add_paragraph(
            f"Subgroup analysis 按年龄组、性别、入院类型和既往住院次数汇总最终模型表现，共生成 {len(subgroup)} 行分层结果。"
            "该分析用于发现模型表现异质性和后续公平性分析线索，不解释为因果效应或临床处置建议。"
        )
        add_caption(document, "Table 8. Subgroup performance summary.")
        add_dataframe_table(
            document,
            subgroup_summary,
            list(subgroup_summary.columns),
            width_dxa=9360,
            font_size=7.2,
        )

    if main_section_enabled(config, "next_writing_suggestions"):
        document.add_heading("下一步写作建议", level=1)
        add_note_table(
            document,
            [
                ("主线一", "把 prediction time 作为核心方法学变量，报告入院时、24h、出院前三个场景。"),
                ("主线二", "把 leakage audit 作为质量控制，不把 Agent 包装成医学问答工具。"),
                ("主线三", "把 selected feature set 作为可解释、低复杂度的补充分析，而不是声称发现因果风险因素。"),
                ("补充材料", f"{configured_supplement_label(config)} 已按配置导出到补充附录 DOCX 和对应 CSV。"),
            ],
        )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def add_supplement_document(project_root: Path, output: Path, config: dict[str, Any]) -> None:
    document = Document()
    apply_compact_styles(document)
    configure_page(document, landscape=True)
    add_footer(document, "ChronoEHR-Agent | Supplementary Appendix")

    title = document.add_paragraph(style="Title")
    title.add_run("Supplementary Appendix")
    document.add_paragraph(
        "本补充材料由 ChronoEHR-Agent 自动生成，用于支持慢病 EHR time-aware prediction benchmark。"
        "它只总结本地 EHR 数据研究结果，不提供医学诊疗建议。"
    )
    add_note_table(
        document,
        [
            ("Design preset", "compact_reference_guide；命名覆盖：wide_table_landscape，US Letter landscape，0.55 inch margins，表格固定宽度 14256 DXA。"),
            ("CSV source", "outputs/tables/supplementary_appendix/"),
            ("Use", "可作为论文补充表初稿；正式投稿前需按目标期刊格式重新整理。"),
        ],
        width_dxa=14256,
    )

    specs = [
        (
            "S1",
            "Table S1. Cohort Summary",
            "table_s1_cohort_summary.csv",
            None,
            8.0,
        ),
        (
            "S2",
            "Table S2. Model Baseline Summary",
            "table_s2_model_baselines.csv",
            None,
            6.2,
        ),
        (
            "S3",
            "Table S3. Feature Group Ablation",
            "table_s3_feature_group_ablation.csv",
            None,
            6.8,
        ),
        (
            "S4",
            "Table S4. Repeated Feature Concepts",
            "table_s4_repeated_feature_concepts.csv",
            30,
            6.5,
        ),
        (
            "S5",
            "Table S5. Full Vs Selected Feature Sets",
            "table_s5_selected_feature_set_comparison.csv",
            None,
            5.8,
        ),
        (
            "S6",
            "Table S6. Leakage Gate Warnings",
            "table_s6_leakage_gate_warnings.csv",
            None,
            6.0,
        ),
        (
            "S7",
            "Table S7. ED Length-of-Stay Sensitivity",
            "table_s7_ed_los_sensitivity.csv",
            None,
            5.6,
        ),
        (
            "S8",
            "Table S8. Threshold And Alert-Burden Analysis",
            "table_s8_threshold_analysis.csv",
            None,
            5.8,
        ),
        (
            "S9",
            "Table S9. Decision-Curve Net Benefit",
            "table_s9_decision_curve.csv",
            None,
            5.5,
        ),
        (
            "S10",
            "Table S10. Subgroup Performance",
            "table_s10_subgroup_performance.csv",
            90,
            4.7,
        ),
    ]
    note_rows = [
        ("S1", "Supports cohort description."),
        ("S2", "Supports traditional baseline comparison."),
        ("S3", "Supports feature group ablation and prediction-time feature availability claims."),
        ("S4", "Supports selected feature set construction."),
        ("S5", "Supports sensitivity analysis showing selected feature sets retain most performance."),
        ("S6", "Documents non-critical leakage-gate warnings, mainly conditional review for ed_los_hours at 24h prediction."),
        ("S7", "Supports the sensitivity analysis showing 24h models are not materially dependent on ed_los_hours."),
        ("S8", "Supports fixed alert-burden interpretation for final 24h and discharge models."),
        ("S9", "Supports decision-curve net-benefit interpretation across common threshold probabilities."),
        ("S10", "Supports subgroup performance summaries and future heterogeneity/fairness checks."),
    ]
    included_note_rows = []
    for table_id, title_text, filename, max_rows, font_size in specs:
        if not supplement_table_enabled(config, table_id):
            continue
        df = read_csv(project_root, SUPPLEMENT_DIR / filename)
        if df.empty:
            continue
        effective_max_rows = configured_max_rows(config, table_id, max_rows)
        document.add_heading(title_text, level=1)
        add_dataframe_table(
            document,
            df,
            list(df.columns),
            width_dxa=14256,
            max_rows=effective_max_rows,
            font_size=font_size,
        )
        add_caption(document, f"Source CSV: outputs/tables/supplementary_appendix/{filename}")
        included_note_rows.extend([row for row in note_rows if row[0] == table_id])

    document.add_heading("Notes For Manuscript Use", level=1)
    add_note_table(
        document,
        included_note_rows,
        width_dxa=14256,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project-root", type=Path, default=DEFAULT_PROJECT)
    parser.add_argument(
        "--export-config",
        type=Path,
        help="Optional JSON manuscript export config. See configs/manuscript_export_template.json.",
    )
    parser.add_argument(
        "--main-output",
        type=Path,
        default=None,
        help="Output path for the main manuscript draft DOCX, relative to project root unless absolute.",
    )
    parser.add_argument(
        "--supplement-output",
        type=Path,
        default=None,
        help="Output path for the supplementary appendix DOCX, relative to project root unless absolute.",
    )
    return parser.parse_args()


def resolve_output(project_root: Path, output: Path) -> Path:
    return output if output.is_absolute() else project_root / output


def main() -> None:
    args = parse_args()
    config_path = None
    if args.export_config is not None:
        config_path = args.export_config if args.export_config.is_absolute() else args.project_root / args.export_config
    config = load_export_config(config_path)
    configured_outputs = config.get("outputs", {})
    main_output_arg = args.main_output or Path(configured_outputs.get("main_docx", DEFAULT_EXPORT_CONFIG["outputs"]["main_docx"]))
    supplement_output_arg = args.supplement_output or Path(
        configured_outputs.get("supplement_docx", DEFAULT_EXPORT_CONFIG["outputs"]["supplement_docx"])
    )
    main_output = resolve_output(args.project_root, main_output_arg)
    supplement_output = resolve_output(args.project_root, supplement_output_arg)
    add_main_document(args.project_root, main_output, config)
    add_supplement_document(args.project_root, supplement_output, config)
    print(f"Wrote {main_output}")
    print(f"Wrote {supplement_output}")


if __name__ == "__main__":
    main()
