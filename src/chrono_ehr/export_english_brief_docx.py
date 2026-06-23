#!/usr/bin/env python3
"""Export the reviewed English Markdown brief as a DOCX draft."""

from __future__ import annotations

import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - depends on workspace runtime.
    raise SystemExit(
        "python-docx is required. Run this via the bundled workspace Python or "
        "through `python3 src/chrono_ehr/run_study.py --english-brief-docx`."
    ) from exc

DEFAULT_PROJECT = Path(__file__).resolve().parents[2]

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "E8EEF5"
BORDER = "B8C2CC"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project-root", type=Path, default=DEFAULT_PROJECT)
    parser.add_argument(
        "--source",
        type=Path,
        default=Path("outputs/reports/chronic_disease_methods_results_english_brief.md"),
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("outputs/reports/ChronoEHR_Methods_Results_English_Brief.docx"),
    )
    return parser.parse_args()


def resolve(project_root: Path, path: Path) -> Path:
    return path if path.is_absolute() else project_root / path


def remove_style_paragraph_borders(style) -> None:
    p_pr = style._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def apply_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(19)
    title.font.color.rgb = DARK_BLUE
    remove_style_paragraph_borders(title)

    for style_name, size, color in [
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12.5, DARK_BLUE),
        ("Heading 3", 11.5, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def add_footer(document: Document) -> None:
    paragraph = document.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = "ChronoEHR-Agent | English Brief Draft | Research use only"
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def clean_inline(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def parse_markdown_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
        rows.pop(1)
    header = rows[0] if rows else []
    body = rows[1:] if len(rows) > 1 else []
    return header, body


def add_markdown_table(document: Document, lines: list[str]) -> None:
    header, body = parse_markdown_table(lines)
    if not header:
        return
    table = document.add_table(rows=1, cols=len(header))
    set_table_borders(table)
    table.autofit = True
    for index, text in enumerate(header):
        cell = table.rows[0].cells[index]
        cell.text = clean_inline(text)
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = DARK_BLUE
    for row in body:
        cells = table.add_row().cells
        for index, text in enumerate(row[: len(header)]):
            cells[index].text = clean_inline(text)
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8)


def add_paragraph(document: Document, text: str) -> None:
    stripped = clean_inline(text.strip())
    if not stripped:
        return
    if stripped.startswith("- "):
        paragraph = document.add_paragraph(stripped[2:], style="List Bullet")
    else:
        paragraph = document.add_paragraph(stripped)
    paragraph.paragraph_format.space_after = Pt(5)


def add_boundary_box(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_borders(table)
    cell = table.rows[0].cells[0]
    cell.text = (
        "Research-use boundary: this document summarizes retrospective EHR data analysis. "
        "It is not a medical QA system, not a clinical decision-support tool, and not a diagnosis or treatment recommendation."
    )
    set_cell_shading(cell, "F4F6F9")
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED


def render_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue
        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.add_run(clean_inline(stripped[2:]))
            add_boundary_box(document)
        elif stripped.startswith("## "):
            document.add_heading(clean_inline(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            document.add_heading(clean_inline(stripped[4:]), level=2)
        else:
            add_paragraph(document, stripped)
        index += 1


def export_docx(source: Path, output: Path) -> None:
    markdown = source.read_text(encoding="utf-8")
    document = Document()
    apply_styles(document)
    configure_page(document)
    add_footer(document)
    render_markdown(document, markdown)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    args = parse_args()
    source = resolve(args.project_root, args.source)
    output = resolve(args.project_root, args.output)
    if not source.exists():
        raise FileNotFoundError(f"Missing English brief Markdown: {source}")
    export_docx(source, output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    main()
